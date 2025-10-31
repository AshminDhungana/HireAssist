from reportlab.pdfgen import canvas
from pypdf import PdfWriter, PdfReader
import io
import os
from docx import Document as DocxDocument

def create_test_resume(tmp_path=None):
    """Create test resume files in the given directory or 'tests/test_data' by default."""
    if tmp_path is None:
        os.makedirs("tests/test_data", exist_ok=True)
        output_dir = "tests/test_data"
    else:
        output_dir = str(tmp_path)

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    
    # Add test content
    c.drawString(100, 750, "Test Resume")
    c.drawString(100, 700, "Work Experience")
    c.drawString(100, 680, "Senior Developer at Tech Corp")
    c.drawString(100, 660, "2020-2023")
    c.drawString(100, 620, "Education")
    c.drawString(100, 600, "BS Computer Science")
    c.drawString(100, 580, "Tech University, 2015-2019")
    c.save()
    
    # Save the regular test resume
    buffer.seek(0)
    test_resume_path = os.path.join(output_dir, "test_resume.pdf")
    with open(test_resume_path, "wb") as f:
        f.write(buffer.getvalue())
    
    # Create encrypted version
    buffer.seek(0)
    reader = PdfReader(buffer)
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.encrypt("test123")
    
    encrypted_path = os.path.join(output_dir, "encrypted.pdf")
    with open(encrypted_path, "wb") as f:
        writer.write(f)

    # Create a DOCX version
    try:
        doc = DocxDocument()
        doc.add_paragraph("Test Resume")
        doc.add_paragraph("Work Experience")
        doc.add_paragraph("Senior Developer at Tech Corp")
        doc.add_paragraph("2020-2023")
        doc.add_paragraph("Education")
        doc.add_paragraph("BS Computer Science")
        doc.add_paragraph("Tech University, 2015-2019")
        docx_path = os.path.join(output_dir, "test_resume.docx")
        doc.save(docx_path)
    except Exception:
        # If python-docx isn't available, skip creating DOCX for tests
        pass

if __name__ == "__main__":
    create_test_resume()